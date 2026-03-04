"""DocxExtractor — extracts text from DOCX files using python-docx."""

from __future__ import annotations

from typing import TYPE_CHECKING, ClassVar

from ..core.base import BaseExtractor, ExtractionError, ExtractionResult

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.run import Run


class DocxExtractor(BaseExtractor):
    """Extracts text from DOCX files using python-docx."""

    supported_extensions: ClassVar[set[str]] = {".docx"}
    required_packages: ClassVar[set[str]] = {"docx"}

    def extract(self, source: str) -> ExtractionResult:
        try:
            from docx import Document
        except ImportError as e:
            raise ExtractionError(
                "python-docx is required for DOCX extraction: pip install python-docx",
                source=source,
                cause=e,
            )

        try:
            doc = Document(source)
        except Exception as e:
            raise ExtractionError(
                f"Failed to open DOCX {source}: {e}", source=source, cause=e
            )

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        metadata: dict = {}
        core = doc.core_properties
        if core.author:
            metadata["Author"] = core.author
        if core.title:
            metadata["Title"] = core.title
        if core.subject:
            metadata["Subject"] = core.subject

        markdown_text = self._to_markdown(doc)

        return ExtractionResult(
            text=full_text,
            source=source,
            source_type="docx",
            extractor_name=self.__class__.__name__,
            metadata=metadata,
            markdown_text=markdown_text or None,
        )

    def _to_markdown(self, doc: Document) -> str:
        """Convert document paragraphs to markdown."""
        lines: list[str] = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            # Headings
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                level = min(level, 6)
                lines.append(f"{'#' * level} {para.text}")
            elif style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
                lines.append(f"- {self._runs_to_markdown(para.runs)}")
            elif style_name in ("List Number", "List Number 2", "List Number 3"):
                lines.append(f"1. {self._runs_to_markdown(para.runs)}")
            else:
                lines.append(self._runs_to_markdown(para.runs))
        return "\n\n".join(lines)

    @staticmethod
    def _runs_to_markdown(runs: list[Run]) -> str:
        """Convert runs to markdown with inline formatting."""
        parts: list[str] = []
        for run in runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            parts.append(text)
        return "".join(parts)
